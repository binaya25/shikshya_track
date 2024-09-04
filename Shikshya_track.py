import math
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
import pandas as pd
import PyPDF2
from docx import Document
from docx.shared import Inches

# Create the main window
root = tk.Tk()
root.title("Shikshya Learning Academy - Finance Tracker")

# Create the logo label


# Create the Income table
income_label = tk.Label(root, text="Income")
income_label.grid(row=1, column=0, padx=10, pady=5)

income_tree = ttk.Treeview(root, columns=("S.No.", "Item", "Student Name", "Program", "Amount", "Remarks"))
income_tree.grid(row=2, column=0, columnspan=4, padx=10, pady=5)

income_tree.heading("#0", text="S.No.")
income_tree.heading("Item", text="Item")
income_tree.heading("Student Name", text="Student Name")
income_tree.heading("Program", text="Program")
income_tree.heading("Amount", text="Amount")
income_tree.heading("Remarks", text="Remarks")

income_items = ["Fee", "Stationeries", "Rental Activities", "Others"]
for i, item in enumerate(income_items, start=1):
    income_tree.insert("", "end", text=str(i), values=(item, "", "", "", "", ""))

# Create the Expense table
expense_label = tk.Label(root, text="Expense")
expense_label.grid(row=1, column=2, padx=10, pady=5)

expense_tree = ttk.Treeview(root, columns=("S.No.", "Item", "Amount", "Remarks"))
expense_tree.grid(row=2, column=2, columnspan=2, padx=10, pady=5)

expense_tree.heading("#0", text="S.No.")
expense_tree.heading("Item", text="Item")
expense_tree.heading("Amount", text="Amount")
expense_tree.heading("Remarks", text="Remarks")

expense_items = ["Rent", "Utilities (electricity, water)", "Stationeries", "Drinks", "Others"]
for i, item in enumerate(expense_items, start=1):
    expense_tree.insert("", "end", text=str(i), values=(item, "", "", ""))

# Create the download button
def download_data():
    # Get the data from the tables
    income_data = [(income_tree.item(child)["values"]) for child in income_tree.get_children()]
    expense_data = [(expense_tree.item(child)["values"]) for child in expense_tree.get_children()]

    # Create a DataFrame from the data
    income_df = pd.DataFrame(income_data, columns=("Item", "Student Name", "Program", "Amount", "Remarks"))
    expense_df = pd.DataFrame(expense_data, columns=("Item", "Amount", "Remarks"))

    # Ask the user to choose the file format
    file_format = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel file", "*.xlsx"), ("PDF file", "*.pdf"), ("Word document", "*.docx")])

    if file_format:
        if file_format.endswith(".xlsx"):
            income_df.to_excel(file_format, index=False)
            expense_df.to_excel(file_format, index=False, sheet_name="Expense")
        elif file_format.endswith(".pdf"):
            # Create a new PDF file
            pdf = PyPDF2.PdfFileWriter()

            # Create a new page for the income data
            income_page = pdf.addPage()
            income_page.mergePage(pdf.getPage(0))  # Merge with the first page
            income_page.mergePage(pdf.getPage(0))  # Merge with the first page
            income_table = [[item for item in row] for row in income_data]
            income_table.insert(0, income_df.columns)
            self._write_table_to_pdf(income_page, income_table)

            # Create a new page for the expense data
            expense_page = pdf.addPage()
            expense_page.mergePage(pdf.getPage(0))  # Merge with the first page
            expense_page.mergePage(pdf.getPage(0))  # Merge with the first page
            expense_table = [[item for item in row] for row in expense_data]
            expense_table.insert(0, expense_df.columns)
            self._write_table_to_pdf(expense_page, expense_table)

            # Write the PDF file
            with open(file_format, "wb") as f:
                pdf.write(f)
        elif file_format.endswith(".docx"):
            document = Document()
            document.add_heading("Income", 0)
            income_table = document.add_table(rows=1, cols=len(income_df.columns))
            for row in income_data:
                row_cells = income_table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
            document.add_heading("Expense", 0)
            expense_table = document.add_table(rows=1, cols=len(expense_df.columns))
            for row in expense_data:
                row_cells = expense_table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
            document.save(file_format)

def _write_table_to_pdf(self, page, table):
    # Write the table to the PDF page
    pass

# Create the input fields
income_items_var = [tk.StringVar() for _ in income_items]
expense_items_var = [tk.StringVar() for _ in expense_items]

for i, item in enumerate(income_items):
    income_item_entry = tk.Entry(root, textvariable=income_items_var[i])
    income_item_entry.grid(row=i+3, column=0)

for i, item in enumerate(expense_items):
    expense_item_entry = tk.Entry(root, textvariable=expense_items_var[i])
    expense_item_entry.grid(row=i+3, column=2)

download_button = tk.Button(root, text="Download", command=download_data)
download_button.grid(row=len(income_items)+len(expense_items)+3, column=2, padx=10, pady=10)

# Run the main loop
root.mainloop()